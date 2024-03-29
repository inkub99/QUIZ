from __future__ import annotations
import streamlit as st

from src.quiz.get_quiz_questions import get_quiz
from src.quiz.get_quiz_questions import load_quiz
from src.utils import config
import docx
from docx import Document
from docx.shared import Pt
import comtypes
import comtypes.client
import os
os.environ["PY_COM_TYPES_DEBUG"] = "1"


def replace_text_in_runs(runs, old_text, new_text):
    for run in runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            new_run = run._element
            for elem in new_run:
                if elem.tag.endswith('rPr'):
                    new_run.insert(0, elem)

def replace_text_in_docx(doc, old_text, new_text):
    for paragraph in doc.paragraphs:
        replace_text_in_runs(paragraph.runs, old_text, new_text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_runs(cell.paragraphs[0].runs, old_text, new_text)


def display_question():
    # Pierwsze pytanie
    if len(st.session_state.questions) == 0:
        quiz = load_quiz()
        first_question = quiz[0]
        st.session_state.questions.append(first_question)

    # Disable the submit button if the user has already answered this question
    submit_button_disabled = (
        st.session_state.current_question in st.session_state.answers
    )

    # Get the current question from the questions list
    try:
        question = st.session_state.questions[st.session_state.current_question]
        # Display the question prompt
        st.write(f"{st.session_state.current_question + 1}. {question['question']}")
        
        # Use an empty placeholder to display the radio button options
        options = st.empty()

        # Display the radio button options and wait for the user to select an answer
        user_answer = options.radio(
        config.config()["app"]["quiz"]["choice"],
        question["options"],
        key=st.session_state.current_question,
        index=None)

        # Display the submit button and disable it if necessary
        submit_button = st.button(
            config.config()["app"]["quiz"]["submit"],
            disabled=submit_button_disabled,)

        # If the user has already answered this question, display their previous answer
        if st.session_state.current_question in st.session_state.answers:
            index = st.session_state.answers[st.session_state.current_question]
            options.radio(
                config.config()["app"]["quiz"]["choice"],
                question["options"],
                key=float(st.session_state.current_question),
                index=index)

        # If the user clicks the submit button, check their answer and show the explanation
        if submit_button:
            # Record the user's answer in the session state
            st.session_state.answers[st.session_state.current_question] = question[
                "options"
            ].index(user_answer)

            # Check if the user's answer is correct and update the score
            if user_answer == question["answer"]:
                st.write(config.config()["app"]["quiz"]["correct"])
                st.session_state.right_answers += 1
            else:
                st.write(f"{config.config()['app']['quiz']['wrong']}{question['answer']}.")
                st.session_state.wrong_answers += 1

            # Show an expander with the explanation of the correct answer
            with st.expander(config.config()["app"]["quiz"]["explanation"]):
                st.write(question["explanation"])

        # Display the current score
        st.write(
            f"{config.config()['app']['quiz']['counter_right']}{st.session_state.right_answers}")
        st.write(
            f"{config.config()['app']['quiz']['counter_wrong']}{st.session_state.wrong_answers}")

    except IndexError:
        st.markdown(config.config()["app"]["quiz"]["over_message"])
        st.markdown(config.config()["app"]["quiz"]["score_message"])
        # Display the current score
        st.write(
            f"{config.config()['app']['quiz']['counter_right']}{st.session_state.right_answers}",
        )
        st.write(
            f"{config.config()['app']['quiz']['counter_wrong']}{st.session_state.wrong_answers}",
        )
        file_path = "PBC_certyfikat_wzor.docx"

        def convert_docx_to_pdf(docx_path, pdf_path):
            pdfkit.from_file(docx_path, pdf_path)
        
        def download_report():
            file_path = "PBC_certyfikat_wzor.docx"
            doc = Document(file_path)
            replace_text_in_docx(doc, "Jan Kowalski", st.session_state.name)
            if str(st.session_state.name).split(' ')[0][-1] == 'a' or str(st.session_state.name).split(' ')[0][-1] == 'A':
                replace_text_in_docx(doc, "Ukończył", "Ukończyła")
    
            # Zapisz plik DOCX
            docx_file_path = "PBC_certyfikat_2.docx"
            doc.save(docx_file_path)
            
            pdf_file_path = f"{str(st.session_state.name)}_PBC_certyfikat.pdf"

            doc = docx.Document(docx_file_path)

            word = comtypes.client.CreateObject("Word.Application")
            docx_path = os.path.abspath(docx_file_path)
            pdf_path = os.path.abspath(pdf_file_path)

            pdf_format = 17  # PDF file format code
            word.Visible = False
            in_file = word.Documents.Open(docx_path)
            in_file.SaveAs(pdf_path, FileFormat=pdf_format)
            in_file.Close()

            word.Quit()
        

            # Odczytaj plik PDF
            with open(pdf_file_path, "rb") as f:
                pdf_bytes = f.read()

            return pdf_bytes

        try:
            if st.session_state.right_answers > 3 and len(str(st.session_state.name))>3:
                st.download_button(
                label="Pobierz dyplom",
                data =download_report(),
                file_name=f"{str(st.session_state.name)}_PBC_certyfikat.pdf",
                mime="application/pdf"
            )
        except:
            pass
                
     

def next_question():
    st.session_state.current_question += 1
    quiz = load_quiz()
    if st.session_state.current_question > len(st.session_state.questions) - 1:
        try:
            next_question = quiz[st.session_state.current_question]
            st.session_state.questions.append(next_question)
        except:
            pass
    
# Define a function to go to the previous question
def prev_question():
    if st.session_state.current_question > 0:
        st.session_state.current_question -= 1
        st.session_state.explanation = None